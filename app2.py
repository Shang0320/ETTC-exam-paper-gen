import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
import random
import io
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2 import service_account

# 主題設定
st.set_page_config(page_title="試卷生成器", page_icon="📄", layout="wide")

# Google Drive 資料夾 ID
ROOT_FOLDER_ID = '17Bcgo8ZeHz0yVhfIxBk7L2wzoiZcyoXt'

# 建立 Google Drive API 服務
def create_drive_service():
    service_account_info = st.secrets["service_account_json"]
    credentials = service_account.Credentials.from_service_account_info(
        service_account_info,
        scopes=['https://www.googleapis.com/auth/drive.readonly']
    )
    return build('drive', 'v3', credentials=credentials)

# 遞迴列出指定資料夾內的檔案
def list_files(service, folder_id):
    query = f"'{folder_id}' in parents and trashed=false"
    result = service.files().list(q=query, fields='files(id, name, mimeType)').execute()
    return result.get('files', [])

# 下載檔案為二進位格式
def download_file(service, file_id):
    request = service.files().get_media(fileId=file_id)
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    fh.seek(0)
    return fh

# 生成試卷
def generate_exam(selected_files, service, class_name, exam_type, subject):
    exam_papers = {}

    for paper_type in ["A卷", "B卷"]:
        doc = Document()

        # 設置頁面大小與邊距
        section = doc.sections[-1]
        section.page_height, section.page_width = Cm(42.0), Cm(29.7)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.top_margin = section.bottom_margin = Cm(1.5 / 2.54)
        section.left_margin = section.right_margin = Cm(2 / 2.54)

        # 添加標題
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(f"海巡署教育訓練測考中心{class_name}梯志願士兵司法警察專長班{exam_type}測驗階段考試（{subject}{paper_type}）")
        header_run.font.name = '標楷體'
        header_run.font.size = Pt(20)
        header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加考試信息
        exam_info_para = doc.add_paragraph("選擇題：100％（共50題，每題2分）")
        exam_info_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in exam_info_para.runs:
            run.font.name = '標楷體'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            run.font.size = Pt(16)

        question_number = 1
        difficulty_counts = {'難': 0, '中': 0, '易': 0}
        total_questions = 0  # 用於計算總題目數

        for file_id in selected_files:
            file_content = download_file(service, file_id)
            df = pd.read_excel(file_content, engine='openpyxl')
            random.seed(1 if paper_type == "A卷" else 2)

            # 確保每題庫抽取的題目不超過 10 題，並且總題數不超過 50 題
            remaining_questions = 50 - total_questions
            if remaining_questions <= 0:
                break

            selected_rows = df.sample(n=min(10, len(df), remaining_questions))

            for _, row in selected_rows.iterrows():
                difficulty_counts['難' if '（難）' in row.iloc[1] else '中' if '（中）' in row.iloc[1] else '易'] += 1
                question_text = f"（{row.iloc[0]}）{question_number}、{row.iloc[1]}"
                question_para = doc.add_paragraph(question_text)

                # 段落格式設置
                paragraph_format = question_para.paragraph_format
                paragraph_format.left_indent = Cm(0)
                paragraph_format.right_indent = Cm(0)
                paragraph_format.hanging_indent = Pt(4 * 0.35)
                paragraph_format.space_after = Pt(0)
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                for run in question_para.runs:
                    run.font.name = '標楷體'
                    run.font.size = Pt(16)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

                question_number += 1
                total_questions += 1

        # 添加難度統計
        summary_text = f"難：{difficulty_counts['難']}，中：{difficulty_counts['中']}，易：{difficulty_counts['易']}"
        summary_para = doc.add_paragraph(summary_text)
        summary_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        exam_papers[paper_type] = buffer.getvalue()

    return exam_papers

# 主程式
service = create_drive_service()
st.markdown("## 📋 基本設定")
class_name = st.text_input("班級名稱", value="113-X", help="請輸入班級名稱，例如：113-1")
exam_type = st.selectbox("考試類型", ["期中", "期末"], help="選擇期中或期末考試")
subject = st.selectbox("科目", ["請選擇", "法律", "專業"], help="選擇科目類型")

if subject and subject != "請選擇":
    st.markdown(f"### 已選科目：{subject}")
    folders = list_files(service, ROOT_FOLDER_ID)
    subject_folder = next((folder for folder in folders if folder['name'] == subject), None)

    if subject_folder:
        topic_files = list_files(service, subject_folder['id'])
        topic_options = {file['name']: file['id'] for file in topic_files if file['mimeType'] == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'}

        if topic_options:
            selected_files = st.multiselect("選擇題庫檔案（限 6 個）", options=list(topic_options.keys()))

            if len(selected_files) == 6 and st.button("生成考卷"):
                selected_file_ids = [topic_options[name] for name in selected_files]
                st.info("正在生成試卷，請稍候...")
                exam_papers = generate_exam(selected_file_ids, service, class_name, exam_type, subject)
                st.success("試卷生成完成！")

                for paper_type, file_data in exam_papers.items():
                    st.download_button(
                        label=f"下載 {paper_type}",
                        data=file_data,
                        file_name=f"{class_name}_{exam_type}_{subject}_{paper_type}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
        else:
            st.warning("未找到任何題庫檔案，請確認資料夾內容！")
    else:
        st.error("未找到對應的科目資料夾，請確認設置！")
