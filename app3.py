import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
import random
import io
import time

# 主題設定
st.set_page_config(page_title="試卷生成器", page_icon="📄", layout="wide")

# 頁面標題與簡介
st.markdown("""
# 📄 志兵班試卷生成器WEB UI
**輕鬆生成專業格式的試卷！**  
請依下列步驟完成試卷生成：
1. 填寫基本資訊。
2. 上傳題庫檔案（6 個 Excel 文件）。
3. 點擊生成按鈕，下載 A 卷與 B 卷試卷。
4. 題庫下載點－ https://drive.google.com/drive/folders/17Bcgo8ZeHz0yVhfIxBk7L2wzoiZcyoXt?usp=sharing
""")

# 分隔線
st.divider()

# 主體內容佈局
col1, col2 = st.columns([1, 2])

with col1:
    st.markdown("## 📋 基本設定")
    class_name = st.text_input("班級名稱", value="113-X", help="請輸入班級名稱，例如：113-1")
    exam_type = st.selectbox("考試類型", ["期中", "期末"], help="選擇期中或期末考試")
    subject = st.selectbox("科目", ["法律", "專業"], help="選擇科目類型")
    
    # 新增是否顯示答案選項
    show_answers = st.checkbox("在試卷上顯示答案", value=False, help="選擇是否在生成的試卷上顯示正確答案")
    
    # 是否包含必考題
    include_required = st.checkbox("優先包含必考題", value=True, help="選擇是否優先選擇必考題目")

with col2:
    st.markdown("## 📤 上傳題庫")
    st.markdown("請上傳 **6 個 Excel 文件**，每個文件代表一個題庫")
    uploaded_files = st.file_uploader("上傳題庫檔案（最多 6 個）", accept_multiple_files=True, type=["xlsx"])

if uploaded_files:
    st.success(f"✅ 已成功上傳 {len(uploaded_files)} 個檔案！")
    if len(uploaded_files) != 6:
        st.warning("⚠️ 請上傳 6 個文件，否則無法生成完整試卷。")
    else:
        st.subheader("檢查上傳文件頭部（前5行）")
        for i, file in enumerate(uploaded_files):
            df = pd.read_excel(file)
            st.write(f"檔案 {i+1} 的頭部：")
            st.write(df.head())

# 初始化 Session State 中的緩存
if "exam_papers" not in st.session_state:
    st.session_state.exam_papers = {}

# 建立一個全域列表，用來記錄各題庫中 A 卷已抽取題目的原始索引（不重複出題）
used_indices = [set() for _ in range(len(uploaded_files))]

# 分隔線
st.divider()

if uploaded_files and len(uploaded_files) == 6:
    if st.button("✨ 開始生成試卷"):
        start_time = time.time()  # 記錄開始時間

        # 各題庫總抽題分配（總題數 50 題）
        total_distribution = [9, 9, 8, 8, 8, 8]

        # A 卷較偏難，設定較高難題數分配（例如：[4,3,3,3,3,3]，總和 19 題）
        A_hard_distribution = [4, 3, 3, 3, 3, 3]
        # B 卷較偏易，設定較低難題數分配（例如：[2,2,2,2,2,2]，總和 12 題）
        B_hard_distribution = [2, 2, 2, 2, 2, 2]

        # 定義生成試卷的函式
        def generate_exam(paper_type, total_distribution, hard_distribution):
            doc = Document()

            # 設置頁面大小與邊距
            section = doc.sections[-1]
            section.page_height, section.page_width = Cm(42.0), Cm(29.7)
            section.orientation = WD_ORIENT.LANDSCAPE
            section.top_margin = section.bottom_margin = Cm(1.5 / 2.54)
            section.left_margin = section.right_margin = Cm(2 / 2.54)

            # 添加標題
            header_para = doc.add_paragraph()
            header_run = header_para.add_run(f"海巡署教育訓練測考中心{class_name}梯志願士兵司法警察專長班{exam_type}測驗階段考試（{subject}{paper_type}）")
            header_run.font.name = '標楷體'
            header_run.font.size = Pt(20)
            header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 添加考試信息
            exam_info_para = doc.add_paragraph("選擇題：100％（共50題，每題2分）")
            exam_info_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in exam_info_para.runs:
                run.font.name = '標楷體'
                run.font.size = Pt(16)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

            question_number = 1  # 全卷題號起始值
            # 初始化難度統計
            difficulty_counts = {'難': 0, '中': 0, '易': 0}

            # 正確答案列表（用於生成答案卷）
            answer_key = []

            # 逐一處理每個題庫
            for i, file in enumerate(uploaded_files):
                # 為確保讀取完整檔案，重設檔案指標
                file.seek(0)
                
                try:
                    # 讀取 Excel 文件
                    df = pd.read_excel(file)
                    
                    # 檢查是否有足夠的列
                    if len(df.columns) < 5:  # 確保至少有題目、答案、難度、必考和選項
                        st.error(f"檔案 {i+1} 的列數不足，請確保題庫格式正確！")
                        return None
                    
                    # 動態匹配並重命名欄位
                    expected_columns = ['序號', '難度', '必考', '答案', '題目', '選項1', '選項2', '選項3', '選項4']
                    current_columns = df.columns.tolist()
                    mapping = {}
                    for expected in expected_columns:
                        for current in current_columns:
                            if expected.lower().strip() in current.lower().strip():
                                mapping[current] = expected
                    
                    # 應用映射重命名
                    df = df.rename(columns=mapping)
                    
                    # 確保必要的欄位存在
                    missing = [col for col in expected_columns if col not in df.columns]
                    if missing:
                        st.error(f"檔案 {i+1} 缺少必要欄位：{missing}")
                        return None
                    
                    # 資料清理
                    df = df.dropna(subset=['題目', '答案'])  # 確保至少有題目和答案
                    
                    # 若為 B 卷，先排除 A 卷已抽取的題目
                    if paper_type == "B卷":
                        df = df[~df.index.isin(used_indices[i])]
                    
                    # 將難度欄位統一化
                    df['難度'] = df['難度'].astype(str).str.strip()
                    df.loc[~df['難度'].isin(['難', '中', '易']), '難度'] = '中'  # 默認為中等難度
                    
                    # 將必考欄位統一化
                    df['必考'] = df['必考'].astype(str).str.strip()
                    df.loc[~df['必考'].isin(['是', '否']), '必考'] = '否'  # 默認為非必考
                    
                    # 確保答案是數字1-4
                    df['答案'] = df['答案'].astype(str).str.strip()
                    df.loc[~df['答案'].isin(['1', '2', '3', '4']), '答案'] = '1'  # 默認答案為1
                    
                    # 如果優先選擇必考題
                    required_questions = pd.DataFrame()
                    if include_required:
                        required_questions = df[df['必考'] == '是']
                        if len(required_questions) == 0:
                            st.warning(f"警告：檔案 {i+1} 沒有必考題，將使用所有題目。")
                        elif len(required_questions) > total_distribution[i]:
                            required_questions = required_questions.sample(n=total_distribution[i], random_state=i+1)
                    
                    # 題庫預處理：先進行隨機排序，保留原始索引
                    seed_shuffle = i + (100 if paper_type == "A卷" else 200)
                    remaining_df = df[~df.index.isin(required_questions.index)]
                    remaining_df = remaining_df.sample(frac=1, random_state=seed_shuffle)

                    total_needed = total_distribution[i] - len(required_questions)
                    desired_hard = hard_distribution[i]
                    random_seed = (1 if paper_type == "A卷" else 2) + i

                    # 根據難度標籤篩選題目
                    df_hard = remaining_df[remaining_df['難度'] == '難']
                    df_medium = remaining_df[remaining_df['難度'] == '中']
                    df_easy = remaining_df[remaining_df['難度'] == '易']

                    if paper_type == "A卷":
                        # A卷：偏向難題
                        req_hard_count = len(required_questions[required_questions['難度'] == '難']) if not required_questions.empty else 0
                        additional_hard_needed = min(desired_hard - req_hard_count, total_needed)
                        additional_hard_needed = max(0, additional_hard_needed)
                        
                        additional_hard = pd.DataFrame()
                        if additional_hard_needed > 0 and len(df_hard) > 0:
                            additional_hard = df_hard.sample(n=min(additional_hard_needed, len(df_hard)), random_state=random_seed)
                        
                        remaining_needed = total_needed - len(additional_hard)
                        remaining_df_for_selection = remaining_df[~remaining_df.index.isin(additional_hard.index)]
                        
                        additional_questions = pd.DataFrame()
                        if remaining_needed > 0 and len(remaining_df_for_selection) > 0:
                            additional_questions = remaining_df_for_selection.sample(n=min(remaining_needed, len(remaining_df_for_selection)), random_state=random_seed)
                        
                        selected_questions = pd.concat([required_questions, additional_hard, additional_questions])
                        selected_questions = selected_questions.sample(frac=1, random_state=random_seed)
                        
                        used_indices[i].update(selected_questions.index.tolist())
                    else:
                        # B卷：偏向易題
                        req_hard_count = len(required_questions[required_questions['難度'] == '難']) if not required_questions.empty else 0
                        additional_hard_needed = min(B_hard_distribution[i] - req_hard_count, total_needed)
                        additional_hard_needed = max(0, additional_hard_needed)
                        
                        additional_hard = pd.DataFrame()
                        if additional_hard_needed > 0 and len(df_hard) > 0:
                            additional_hard = df_hard.sample(n=min(additional_hard_needed, len(df_hard)), random_state=random_seed)
                        
                        remaining_needed = total_needed - len(additional_hard)
                        easy_to_select = min(remaining_needed, len(df_easy))
                        
                        additional_easy = pd.DataFrame()
                        if easy_to_select > 0:
                            additional_easy = df_easy.sample(n=easy_to_select, random_state=random_seed)
                        
                        remaining_needed -= len(additional_easy)
                        remaining_df_for_selection = remaining_df[~remaining_df.index.isin(additional_hard.index) & ~remaining_df.index.isin(additional_easy.index)]
                        
                        additional_questions = pd.DataFrame()
                        if remaining_needed > 0 and len(remaining_df_for_selection) > 0:
                            additional_questions = remaining_df_for_selection.sample(n=min(remaining_needed, len(remaining_df_for_selection)), random_state=random_seed)
                        
                        selected_questions = pd.concat([required_questions, additional_hard, additional_easy, additional_questions])
                        selected_questions = selected_questions.sample(frac=1, random_state=random_seed)

                    # 將抽取的題目依序加入文件，並更新難度統計
                    for _, row in selected_questions.iterrows():
                        answer = row['答案']
                        question_text = row['題目']
                        options = [row['選項1'], row['選項2'], row['選項3'], row['選項4']]
                        difficulty = row['難度']
                        
                        answer_key.append((question_number, answer))
                        
                        options_text = "(".join([f"{i+1}{opt}" for i, opt in enumerate(options)]) + ")"
                        if show_answers:
                            question_para = doc.add_paragraph(f"（{answer}）{question_number}、{question_text}{options_text}（{difficulty}）")
                        else:
                            question_para = doc.add_paragraph(f"{question_number}、{question_text}{options_text}（{difficulty}）")
                        
                        paragraph_format = question_para.paragraph_format
                        paragraph_format.left_indent = Cm(0)
                        paragraph_format.right_indent = Cm(0)
                        paragraph_format.hanging_indent = Pt(8 * 0.35)
                        paragraph_format.space_after = Pt(0)
                        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        for run in question_para.runs:
                            run.font.name = '標楷體'
                            run.font.size = Pt(16)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                        
                        difficulty_counts[difficulty] += 1
                        question_number += 1

                except Exception as e:
                    st.error(f"處理檔案 {i+1} 時發生錯誤: {str(e)}")
                    return None

            # 添加難度統計
            summary_text = f"難：{difficulty_counts['難']}，中：{difficulty_counts['中']}，易：{difficulty_counts['易']}"
            summary_para = doc.add_paragraph(summary_text)
            summary_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # 如果設置顯示答案，添加答案卷
            if not show_answers:
                doc.add_page_break()
                answer_title = doc.add_paragraph()
                answer_title.add_run(f"{subject}{paper_type} 答案卷").bold = True
                answer_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                answers_per_row = 5
                num_rows = (len(answer_key) + answers_per_row - 1) // answers_per_row
                
                for row in range(num_rows):
                    answer_row = doc.add_paragraph()
                    for col in range(answers_per_row):
                        idx = row * answers_per_row + col
                        if idx < len(answer_key):
                            q_num, ans = answer_key[idx]
                            answer_row.add_run(f"{q_num}. {ans}     ")
                    answer_row.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()

        # 分別生成 A 卷與 B 卷
        exam_A = generate_exam("A卷", total_distribution, A_hard_distribution)
        exam_B = generate_exam("B卷", total_distribution, B_hard_distribution)

        if exam_A and exam_B:
            st.session_state.exam_papers["A卷"] = exam_A
            st.session_state.exam_papers["B卷"] = exam_B

            end_time = time.time()
            elapsed_time = end_time - start_time
            st.success(f"🎉 試卷生成完成！耗時：{elapsed_time:.2f} 秒")
        else:
            st.error("❌ 試卷生成失敗，請檢查題庫格式並重試。")

# 顯示下載按鈕
if "exam_papers" in st.session_state and st.session_state.exam_papers:
    st.markdown("## 📥 下載試卷")
    for paper_type, file_data in st.session_state.exam_papers.items():
        st.download_button(
            label=f"下載 {paper_type}",
            data=file_data,
            file_name=f"{class_name}_{exam_type}_{subject}_{paper_type}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
