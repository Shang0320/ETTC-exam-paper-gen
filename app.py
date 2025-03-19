import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
import random
import io
import time

# 主題設定
st.set_page_config(page_title="試卷生成器", page_icon="📄", layout="wide")

# 頁面標題與簡介
st.markdown("""
# 📄 志兵班試卷生成器WEB UI
**輕鬆生成專業格式的試卷！**  
請依下列步驟完成試卷生成：
1. 填寫基本資訊與分配設定。
2. 上傳題庫檔案（6 個 Excel 文件）。
3. 點擊生成按鈕，下載標準化的 A 卷與 B 卷。
4. 題庫下載點－ [點此下載](https://drive.google.com/drive/folders/17Bcgo8ZeHz0yVhfIxBk7L2wzoiZcyoXt?usp=sharing)
""")

# 分隔線
st.divider()

# 主體內容佈局：左側基本設定、右側上傳題庫
col1, col2 = st.columns([1, 2])

with col1:
    st.markdown("## 📋 基本設定")
    class_name = st.text_input("班級名稱", value="113-X", help="請輸入班級名稱，例如：113-1")
    exam_type = st.selectbox("考試類型", ["期中", "期末"], help="選擇期中或期末考試")
    subject = st.selectbox("科目", ["法律", "專業"], help="選擇科目類型")
    
    st.markdown("### A卷 分配設定")
    total_distribution_A_input = st.text_input("請以逗號分隔輸入 A 卷各題庫抽題數（例如：9,9,8,8,8,8）", value="9,9,8,8,8,8")
    hard_distribution_A_input  = st.text_input("請以逗號分隔輸入 A 卷各題庫難題數（上限總數 20，例如：5,2,2,4,3,3）", value="5,2,2,4,3,3")
    
    st.markdown("### B卷 分配設定")
    total_distribution_B_input = st.text_input("請以逗號分隔輸入 B 卷各題庫抽題數（例如：9,9,8,8,8,8）", value="9,9,8,8,8,8")
    hard_distribution_B_input  = st.text_input("請以逗號分隔輸入 B 卷各題庫難題數（上限總數 20，例如：4,3,2,3,4,2）", value="4,3,2,3,4,2")

with col2:
    st.markdown("## 📤 上傳題庫")
    st.markdown("請上傳 **6 個 Excel 文件**，每個文件代表一個題庫")
    uploaded_files = st.file_uploader("上傳題庫檔案（最多 6 個）", accept_multiple_files=True, type=["xlsx"])

if uploaded_files:
    st.success(f"✅ 已成功上傳 {len(uploaded_files)} 個檔案！")
    if len(uploaded_files) != 6:
        st.warning("⚠️ 請上傳 6 個文件，否則無法生成完整試卷。")

# 初始化 Session State 緩存
if "exam_papers" not in st.session_state:
    st.session_state.exam_papers = {}

# 分隔線
st.divider()

if uploaded_files and len(uploaded_files) == 6:
    if st.button("✨ 開始生成試卷"):
        start_time = time.time()  # 記錄開始時間
        
        # 解析各項分配設定
        try:
            total_distribution_A = [int(x.strip()) for x in total_distribution_A_input.split(",")]
            total_distribution_B = [int(x.strip()) for x in total_distribution_B_input.split(",")]
            hard_distribution_A  = [int(x.strip()) for x in hard_distribution_A_input.split(",")]
            hard_distribution_B  = [int(x.strip()) for x in hard_distribution_B_input.split(",")]
            
            if (len(total_distribution_A) != 6 or len(total_distribution_B) != 6 or 
                len(hard_distribution_A) != 6 or len(hard_distribution_B) != 6):
                st.error("請確保所有分配設定均包含 6 個數字。")
                st.stop()
            if sum(total_distribution_A) != 50 or sum(total_distribution_B) != 50:
                st.error("請確保 A 卷與 B 卷抽題分配的總題數均為 50 題。")
                st.stop()
            if sum(hard_distribution_A) > 20 or sum(hard_distribution_B) > 20:
                st.error("每張考卷的難題總數不得超過 20 題。")
                st.stop()
        except Exception as e:
            st.error("分配設定格式錯誤，請檢查輸入。")
            st.stop()
        
        # 用來記錄 A 卷每個題庫已抽題的索引，避免 B 卷重複抽題
        used_indices = [set() for _ in range(6)]
        
        # 定義生成試卷的函式，依據卷別採用不同分配設定
        def generate_exam(paper_type, total_distribution, hard_distribution, exclude_used=False):
            doc = Document()
            # 設置頁面大小與邊距
            section = doc.sections[-1]
            section.page_height, section.page_width = Cm(42.0), Cm(29.7)
            section.orientation = WD_ORIENT.LANDSCAPE
            section.top_margin = section.bottom_margin = Cm(1.5 / 2.54)
            section.left_margin = section.right_margin = Cm(2 / 2.54)
            
            # 添加標題
            header_para = doc.add_paragraph()
            header_run = header_para.add_run(f"海巡署教育訓練測考中心{class_name}梯志願士兵司法警察專長班{exam_type}測驗階段考試（{subject}{paper_type}）")
            header_run.font.name = '標楷體'
            header_run.font.size = Pt(20)
            header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加考試資訊
            exam_info_para = doc.add_paragraph("選擇題：100％（共50題，每題2分）")
            exam_info_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in exam_info_para.runs:
                run.font.name = '標楷體'
                run.font.size = Pt(16)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            
            question_number = 1
            difficulty_counts = {'難': 0, '中': 0, '易': 0}
            
            # 逐一處理每個題庫
            for i, file in enumerate(uploaded_files):
                # 讀取題庫
                df = pd.read_excel(file)
                # 若要排除已被 A 卷抽取的題目，則過濾掉該題庫已標記索引
                if exclude_used:
                    df = df[~df.index.isin(used_indices[i])]
                
                total_needed = total_distribution[i]
                desired_hard = hard_distribution[i]
                random_seed = (1 if paper_type == "A卷" else 2) + i
                
                # 抽取難題（根據題目文字判斷包含「（難）」）
                df_hard = df[df.iloc[:, 1].str.contains('（難）', na=False)]
                n_hard_available = len(df_hard)
                n_hard_to_select = min(desired_hard, total_needed, n_hard_available)
                if n_hard_to_select > 0:
                    selected_hard = df_hard.sample(n=n_hard_to_select, random_state=random_seed)
                else:
                    selected_hard = pd.DataFrame(columns=df.columns)
                
                # 從非難題中抽取剩餘題數
                remaining = total_needed - n_hard_to_select
                df_nonhard = df[~df.index.isin(df_hard.index)]
                n_nonhard_available = len(df_nonhard)
                n_nonhard_to_select = min(remaining, n_nonhard_available)
                if n_nonhard_to_select > 0:
                    selected_nonhard = df_nonhard.sample(n=n_nonhard_to_select, random_state=random_seed)
                else:
                    selected_nonhard = pd.DataFrame(columns=df.columns)
                
                # 合併難題與非難題並隨機排序
                selected_questions = pd.concat([selected_hard, selected_nonhard])
                selected_questions = selected_questions.sample(frac=1, random_state=random_seed).reset_index(drop=True)
                
                # 對 A 卷，記錄該題庫中抽取的題目索引，避免 B 卷重複抽題
                if paper_type == "A卷":
                    used_indices[i].update(selected_hard.index.tolist())
                    used_indices[i].update(selected_nonhard.index.tolist())
                
                # 將抽取的題目依序加入文件
                for _, row in selected_questions.iterrows():
                    question_text = f"（{row.iloc[0]}）{question_number}、{row.iloc[1]}"
                    question_para = doc.add_paragraph(question_text)
                    paragraph_format = question_para.paragraph_format
                    paragraph_format.left_indent = Cm(0)
                    paragraph_format.right_indent = Cm(0)
                    paragraph_format.hanging_indent = Pt(8 * 0.35)
                    paragraph_format.space_after = Pt(0)
                    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    
                    for run in question_para.runs:
                        run.font.name = '標楷體'
                        run.font.size = Pt(16)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                    
                    # 更新難度統計：若題目中包含「（難）」則計入難題，否則依據內容判斷中或易題
                    if '（難）' in row.iloc[1]:
                        difficulty_counts['難'] += 1
                    elif '（中）' in row.iloc[1]:
                        difficulty_counts['中'] += 1
                    else:
                        difficulty_counts['易'] += 1
                        
                    question_number += 1
            
            # 添加難度統計（此段不可刪除）
            summary_text = f"難：{difficulty_counts['難']}，中：{difficulty_counts['中']}，易：{difficulty_counts['易']}"
            summary_para = doc.add_paragraph(summary_text)
            summary_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # 輸出文件至記憶體
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        
        # 生成 A 卷：不排除題目，並標記已抽題
        exam_A = generate_exam("A卷", total_distribution_A, hard_distribution_A, exclude_used=False)
        # 生成 B 卷：排除 A 卷已抽題
        exam_B = generate_exam("B卷", total_distribution_B, hard_distribution_B, exclude_used=True)
        
        # 將生成的試卷存入 Session State
        st.session_state.exam_papers["A卷"] = exam_A
        st.session_state.exam_papers["B卷"] = exam_B
        
        end_time = time.time()  # 記錄結束時間
        elapsed_time = end_time - start_time
        st.success(f"🎉 試卷生成完成！耗時：{elapsed_time:.2f} 秒")

# 顯示下載按鈕
if "exam_papers" in st.session_state and st.session_state.exam_papers:
    st.markdown("## 📥 下載試卷")
    for paper_type, file_data in st.session_state.exam_papers.items():
        st.download_button(
            label=f"下載 {paper_type}",
            data=file_data,
            file_name=f"{class_name}_{exam_type}_{subject}_{paper_type}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
